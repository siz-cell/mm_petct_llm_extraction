# -*- coding: utf-8 -*-
"""Table 2/3/4 重建 — PYCM字段级F1 + SUVmax + Exact Match"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import numpy as np

PROJECT = Path(__file__).resolve().parent.parent
DATA = PROJECT / "data" / "汇总规范化数据.xlsx"
DST = PROJECT / "output" / "tables_final"
DST.mkdir(parents=True, exist_ok=True)

NM = {
    "deepseek-v3-250324":"DeepSeek-V3","deepseek-r1":"DeepSeek-R1","deepseekr1":"DeepSeek-R1",
    "claude-haiku-4-5-20251001":"Claude Haiku 4.5","claude-haiku-4-5-20251001-thinking":"Claude Haiku 4.5 Thinking",
    "gpt-4.1-2025-04-14":"GPT-4.1","o4-mini-2025-04-16-medium":"o4-mini",
    "gemini-2.5-flash-nothinking":"Gemini 2.5 Flash","gemini-2.5-flash-thinking":"Gemini 2.5 Flash Thinking",
    "Qwen3-4B-Instruct-2507":"Qwen3-4B-Instruct","Qwen3-4B-Thinking-2507":"Qwen3-4B-Thinking",
    "qwen3-235b-a22b-Instruct-2507":"Qwen3-235B-Instruct","qwen3-235b-a22b-thinking-2507":"Qwen3-235B-Thinking",
    "qwen3-235b-a22b-Instruct-2507 zero-shot":"Qwen3-235B-Instruct 0-shot",
    "qwen3-235b-a22b-thinking-2507 zero-shot":"Qwen3-235B-Thinking 0-shot",
    "qwen3-235b-a22b-Instruct-2507 8-shot":"Qwen3-235B-Instruct 8-shot",
    "qwen3-235b-a22b-thinking-2507 8-shot":"Qwen3-235B-Thinking 8-shot",
    "Qwen3-4B-Thinking-2507-Lora-Structured 8-shot":"Qwen3-4B-LoRA 8-shot",
    "Qwen3-4B-Thinking-2507-Lora-Structured zero-shot":"Qwen3-4B-LoRA 0-shot",
    "Qwen3-4B-Thinking-2507 8-shot":"Qwen3-4B-Thinking 8-shot",
    "Qwen3-4B-Thinking-2507 zero-shot":"Qwen3-4B-Thinking 0-shot",
}

def en(raw):
    best, bl = raw, 0
    for k,v in NM.items():
        if k in raw and len(k)>bl: best=v; bl=len(k)
    return best

def parse_ci(s):
    if s is None or pd.isna(s): return None,None,None
    if isinstance(s,(int,float)): return float(s),None,None
    m = re.search(r'([\d.]+)\s*[\(（]\s*([\d.]+)[–\-]\s*([\d.]+)[\)）]',str(s))
    if m: return float(m.group(1)),float(m.group(2)),float(m.group(3))
    m = re.search(r'([\d.]+)',str(s)); return (float(m.group(1)),None,None) if m else (None,None,None)

def fmt(v,lo,hi):
    if v is None: return "-"
    if lo is not None and hi is not None:
        return f"{v:.2f} ({lo:.2f}–{hi:.2f})"
    return f"{v:.2f}"

# Load data
df_pycm = pd.read_excel(DATA, sheet_name="PYCM字段级指标")
df_s2 = pd.read_excel(DATA, sheet_name="模型字段对比_临时对应图表.xlsx_Sheet2")
df_s3 = pd.read_excel(DATA, sheet_name="模型字段对比_临时对应图表.xlsx_Sheet3")
df_s4 = pd.read_excel(DATA, sheet_name="模型字段对比_临时对应图表.xlsx_Sheet4")
df_s4 = df_s4[df_s4.iloc[:,1].notna()]

# 10 columns
COLS = [
    "PET Positivity\n(F1)", "FL Location\n(F1)", "FL Count\n(F1)",
    "EM Location\n(F1)", "EM Count\n(F1)", "PM Count\n(F1)",
    "Fracture\n(F1)", "SUVmax\nMAE", "SUVmax\nRMSE", "Exact\nMatch"
]
FIELD_MAP_PYCM = {
    "PET Positivity\n(F1)": "PET阳性/阴性",
    "FL Location\n(F1)": "FLs-部位",
    "FL Count\n(F1)": "FLs-数量",
    "EM Location\n(F1)": "EM-部位",
    "EM Count\n(F1)": "EM-数量",
    "PM Count\n(F1)": "PM-数量",
    "Fracture\n(F1)": "骨折-有/无",
}

def get_f1(model, field_cn):
    sub = df_pycm[(df_pycm["模型"]==model)&(df_pycm["字段"]==field_cn)]
    # Cross-map Sheet3/4 model names to PYCM names
    if len(sub)==0:
        pycm_name = {
            "qwen3-235b-a22b-Instruct-2507 zero-shot":"qwen3-235b-a22b-Instruct-2507",
            "qwen3-235b-a22b-thinking-2507 zero-shot":"qwen3-235b-a22b-thinking-2507",
            "qwen3-235b-a22b-Instruct-2507 8-shot":"qwen3-235b-a22b-instruct-25078shot",
            "qwen3-235b-a22b-thinking-2507 8-shot":"qwen3-235b-a22b-thinking-25078shot",
            "Qwen3-4B-Thinking-2507-Lora-Structured 8-shot":"lora 4b 70 8shot",
            "Qwen3-4B-Thinking-2507-Lora-Structured zero-shot":"lora 4b 70 0shot",
            "Qwen3-4B-Thinking-2507 8-shot":"4b 70 8shot",
            "Qwen3-4B-Thinking-2507 zero-shot":"4b 70 0 shot",
            "deepseek-r1":"deepseekr1",
        }.get(model, model)
        sub = df_pycm[(df_pycm["模型"]==pycm_name)&(df_pycm["字段"]==field_cn)]
    if len(sub):
        r = sub.iloc[0]
        return r["F1"], r["F1_lo"], r["F1_hi"]
    return None,None,None

def build_row(model_raw, suv_df):
    model_en = en(model_raw)
    row = [model_en]
    # F1 fields from PYCM
    for col_name in COLS[:7]:
        field_cn = FIELD_MAP_PYCM[col_name]
        v,lo,hi = get_f1(model_raw, field_cn)
        row.append(fmt(v,lo,hi))
    # SUVmax + Exact Match from structured table
    suv_row = suv_df[suv_df.iloc[:,0]==model_raw]
    if len(suv_row):
        r = suv_row.iloc[0]
        for j in [6,7,8]:
            v,lo,hi = parse_ci(r.iloc[j])
            row.append(fmt(v,lo,hi))
    else:
        row += ["-"]*3
    return row

def make_doc(title):
    d=Document(); s=d.styles["Normal"]; s.font.name="Helvetica"; s.font.size=Pt(9)
    s.element.rPr.rFonts.set(qn("w:eastAsia"),"Helvetica")
    cap=d.add_paragraph()
    for t,b in [("Table X | ",True),(title,True)]:
        r=cap.add_run(t); r.font.name="Helvetica"; r.font.size=Pt(10); r.font.bold=b
    cap.alignment=WD_ALIGN_PARAGRAPH.LEFT; cap.paragraph_format.space_after=Pt(4)
    return d

def add_borders(table):
    tbl=table._tbl; tblBorders=OxmlElement("w:tblBorders")
    def _b(sz):
        b=OxmlElement("w:single"); b.set(qn("w:val"),"single")
        b.set(qn("w:sz"),str(sz)); b.set(qn("w:space"),"0"); b.set(qn("w:color"),"000000"); return b
    tblBorders.append(_b(30)); tblBorders.append(_b(30))
    for s in ["left","right","insideH","insideV"]:
        nil=OxmlElement(f"w:{s}"); nil.set(qn("w:val"),"nil"); tblBorders.append(nil)
    tbl.tblPr.append(tblBorders)
    for cell in table.rows[0].cells:
        tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcBorders=OxmlElement("w:tcBorders")
        for s in ["top","left","right"]:
            nil=OxmlElement(f"w:{s}"); nil.set(qn("w:val"),"nil"); tcBorders.append(nil)
        bot=OxmlElement("w:bottom"); bot.set(qn("w:val"),"single")
        bot.set(qn("w:sz"),"15"); bot.set(qn("w:space"),"0"); bot.set(qn("w:color"),"000000")
        tcBorders.append(bot); tcPr.append(tcBorders)

def save_table(doc, rows, filename, bold_cond=None):
    table=doc.add_table(rows=1+len(rows), cols=1+len(COLS)); table.style="Table Grid"
    # Header
    c=table.cell(0,0); p=c.paragraphs[0]; p.clear()
    r=p.add_run("Model"); r.font.name="Helvetica"; r.font.size=Pt(7); r.font.bold=True
    c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for j,h in enumerate(COLS):
        c=table.cell(0,j+1); p=c.paragraphs[0]; p.clear()
        r=p.add_run(h); r.font.name="Helvetica"; r.font.size=Pt(6.5); r.font.bold=True
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # Data
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            c=table.cell(i+1,j); p=c.paragraphs[0]; p.clear()
            r=p.add_run(str(v)); r.font.name="Helvetica"; r.font.size=Pt(7)
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            if bold_cond and bold_cond in str(row[0]): r.font.bold=True
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if j>=1 else WD_ALIGN_PARAGRAPH.LEFT
    add_borders(table)
    doc.save(DST/filename); print(f"  {filename}")

# ---- T2: 12 models zero-shot ----
models_t2 = df_s2.iloc[:,0].tolist()
rows_t2 = [build_row(m, df_s2) for m in models_t2]
doc=make_doc("Zero-Shot LLM Structured Extraction Performance on 365 PET/CT Reports")
save_table(doc, rows_t2, "T2_ZeroShot.docx", "235B-Thinking")

# ---- T3: 4 variants few-shot ----
models_t3 = df_s3.iloc[:,0].tolist()
rows_t3 = [build_row(m, df_s3) for m in models_t3]
doc=make_doc("Few-Shot (8-Shot) Prompting Performance — Qwen3-235B")
save_table(doc, rows_t3, "T3_FewShot.docx", "Thinking 8-shot")

# ---- T4: 4 variants optimization ----
models_t4 = df_s4.iloc[:,0].tolist()
rows_t4 = [build_row(m, df_s4) for m in models_t4]
doc=make_doc("Optimization Strategies for Small Models (Qwen3-4B)")
save_table(doc, rows_t4, "T4_Optimization.docx", "LoRA 8-shot")

# ---- T9: External validation structured extraction (lora jbz0) ----
df_t9 = pd.read_excel(DATA, sheet_name="外部验证集_Sheet9_set1")
doc=make_doc("External Validation: Structured Extraction Performance (n=162)")
cols_ev=["Metric","Internal (95% CI)","External (95% CI)","P value"]
table=doc.add_table(rows=1+len(df_t9),cols=4); table.style="Table Grid"
for j,h in enumerate(cols_ev):
    c=table.cell(0,j); p=c.paragraphs[0]; p.clear()
    r=p.add_run(h); r.font.name="Helvetica"; r.font.size=Pt(7); r.font.bold=True
    c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for i,row in df_t9.iterrows():
    for j in range(4):
        c=table.cell(i+1,j); p=c.paragraphs[0]; p.clear()
        v=str(row.iloc[j]) if pd.notna(row.iloc[j]) else ""
        r=p.add_run(v); r.font.name="Helvetica"; r.font.size=Pt(7.5)
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
add_borders(table)
doc.add_paragraph().add_run("lora jbz0 internal (n=69) vs external (n=162).").font.size=Pt(7.5)
doc.save(DST/"T9_External_Structured.docx"); print("T9 done")

# ---- T11: External validation downstream ----
df_t11 = pd.read_excel(DATA, sheet_name="外部验证集_Sheet11")
doc=make_doc("External Validation: Downstream Clinical Tasks (n=162)")
valid=df_t11[df_t11.iloc[:,0].notna()&(df_t11.iloc[:,0]!="Metric")]
table=doc.add_table(rows=1+len(valid),cols=4); table.style="Table Grid"
for j,h in enumerate(cols_ev):
    c=table.cell(0,j); p=c.paragraphs[0]; p.clear()
    r=p.add_run(h); r.font.name="Helvetica"; r.font.size=Pt(7); r.font.bold=True
    c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for i,(_,row) in enumerate(valid.iterrows()):
    for j in range(4):
        c=table.cell(i+1,j); p=c.paragraphs[0]; p.clear()
        v=str(row.iloc[j]) if pd.notna(row.iloc[j]) else ""
        r=p.add_run(v); r.font.name="Helvetica"; r.font.size=Pt(7.5)
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
add_borders(table)
doc.add_paragraph().add_run("lora jbz0 baseline+response+MRD internal vs external.").font.size=Pt(7.5)
doc.save(DST/"T11_External_Downstream.docx"); print("T11 done")

print("Done!")
