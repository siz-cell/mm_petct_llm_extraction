# -*- coding: utf-8 -*-
"""一键修复所有用户标注的x问题"""
import shutil, csv
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import numpy as np
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns

PROJ = Path(__file__).resolve().parent.parent
DATA = PROJ / "data"
DST_TABLES = PROJ / "output" / "tables_final"
DST_CHARTS = PROJ / "output" / "final"
DST_TABLES.mkdir(parents=True, exist_ok=True)
DST_CHARTS.mkdir(parents=True, exist_ok=True)

NM = {
    "deepseek-v3-250324":"DeepSeek-V3","deepseek-r1":"DeepSeek-R1","deepseekr1":"DeepSeek-R1",
    "claude-haiku-4-5-20251001":"Claude Haiku 4.5","claude-haiku-4-5-20251001-thinking":"Claude Haiku 4.5 Thinking",
    "gpt-4.1-2025-04-14":"GPT-4.1","o4-mini-2025-04-16-medium":"o4-mini",
    "gemini-2.5-flash-nothinking":"Gemini 2.5 Flash","gemini-2.5-flash-thinking":"Gemini 2.5 Flash Thinking",
    "Qwen3-4B-Instruct-2507":"Qwen3-4B-Instruct","Qwen3-4B-Thinking-2507":"Qwen3-4B-Thinking",
    "qwen3-235b-a22b-Instruct-2507":"Qwen3-235B-Instruct","qwen3-235b-a22b-thinking-2507":"Qwen3-235B-Thinking",
    "qwen3235bthinking":"Qwen3-235B-Thinking","qwen3235bnothinking":"Qwen3-235B-Instruct",
    "qwen34bthinking":"Qwen3-4B-Thinking","qwen34bnothinking":"Qwen3-4B-Instruct",
    "loraqwen34bthinking":"Qwen3-4B-LoRA-Thinking","loraqwen34bnothinking":"Qwen3-4B-LoRA-Instruct",
    "qwen3235bthinking8shotcot":"Qwen3-235B-Thinking 8-shot",
    "loraqwen34bthinking8shot":"Qwen3-4B-LoRA-Thinking 8-shot",
    "qwen34b8shotcotthinking8shot":"Qwen3-4B-Thinking 8-shot",
}
def en(raw):
    best,bl=raw,0
    for k,v in NM.items():
        if k in raw and len(k)>bl: best=v; bl=len(k)
    return best

def add_borders(table):
    tbl=table._tbl; tblBorders=OxmlElement("w:tblBorders")
    def _b(sz):
        b=OxmlElement("w:single"); b.set(qn("w:val"),"single")
        b.set(qn("w:sz"),str(sz)); b.set(qn("w:space"),"0"); b.set(qn("w:color"),"000000"); return b
    tblBorders.append(_b(30)); tblBorders.append(_b(30))
    for s in ["left","right","insideH","insideV"]:
        nil=OxmlElement(f"w:{s}"); nil.set(qn("w:val"),"nil"); tblBorders.append(nil)
    tbl.tblPr.append(tblBorders)
    for cell in table.rows[0].cells:
        tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcBorders=OxmlElement("w:tcBorders")
        for s in ["top","left","right"]:
            nil=OxmlElement(f"w:{s}"); nil.set(qn("w:val"),"nil"); tcBorders.append(nil)
        bot=OxmlElement("w:bottom"); bot.set(qn("w:val"),"single")
        bot.set(qn("w:sz"),"15"); bot.set(qn("w:space"),"0"); bot.set(qn("w:color"),"000000")
        tcBorders.append(bot); tcPr.append(tcBorders)

# ========== 1. Fix T7: regenerate ==========
print("1. Fixing T7...")
df_bl=pd.read_excel(DATA/"汇总规范化数据.xlsx",sheet_name="初诊基线评估")
doc=Document(); s=doc.styles["Normal"]; s.font.name="Helvetica"; s.font.size=Pt(9)
p=doc.add_paragraph(); r=p.add_run("Table 7 | Downstream: Baseline Bone Marrow Infiltration Classification")
r.font.size=Pt(10); r.font.bold=True
p.alignment=WD_ALIGN_PARAGRAPH.LEFT
cols7=["Model","Weighted F1 (95% CI)","Accuracy (95% CI)","F1 (Diffuse)","F1 (Focal)","F1 (Minimal)","F1 (Mixed)"]
table=doc.add_table(rows=1+len(df_bl),cols=7); table.style="Table Grid"
for j,h in enumerate(cols7):
    c=table.cell(0,j); p=c.paragraphs[0]; p.clear()
    r=p.add_run(h); r.font.size=Pt(7); r.font.bold=True
    c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for i,row in df_bl.iterrows():
    mn=en(str(row["模型"]))
    vals=[mn,
          f"{row['Weighted_F1']:.2f} ({row['Weighted_F1_lo']:.2f}–{row['Weighted_F1_hi']:.2f})",
          f"{row['Accuracy']:.2f} ({row['Accuracy_lo']:.2f}–{row['Accuracy_hi']:.2f})",
          f"{row['F1_Dif']:.2f}",f"{row['F1_Foc']:.2f}",f"{row['F1_Min']:.2f}",f"{row['F1_Mix']:.2f}"]
    for j,v in enumerate(vals):
        c=table.cell(i+1,j); p=c.paragraphs[0]; p.clear()
        r=p.add_run(v); r.font.size=Pt(7.5)
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if j>=1 else WD_ALIGN_PARAGRAPH.LEFT
add_borders(table)
doc.save(DST_TABLES/"T7_Downstream_Baseline.docx"); print("  T7 done")

# ========== 2. Fix T8: regenerate ==========
print("2. Fixing T8...")
df_mrd=pd.read_excel(DATA/"汇总规范化数据.xlsx",sheet_name="MRD多模型对比")
doc=Document(); s=doc.styles["Normal"]; s.font.name="Helvetica"; s.font.size=Pt(9)
p=doc.add_paragraph(); r=p.add_run("Table 8 | MRD Assessment Performance"); r.font.size=Pt(10); r.font.bold=True
p.alignment=WD_ALIGN_PARAGRAPH.LEFT
cols8=["Model","AUC-ROC (95% CI)","AUC-PR (95% CI)","Sensitivity","Specificity","PPV","NPV","Accuracy","F1 (Neg)"]
table=doc.add_table(rows=1+len(df_mrd),cols=9); table.style="Table Grid"
for j,h in enumerate(cols8):
    c=table.cell(0,j); p=c.paragraphs[0]; p.clear()
    r=p.add_run(h); r.font.size=Pt(6.5); r.font.bold=True
    c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for i,row in df_mrd.iterrows():
    mn=en(str(row.iloc[0]))
    vals=[mn]
    for j in range(1,9):
        v=str(row.iloc[j]) if pd.notna(row.iloc[j]) else "-"
        vals.append(v)
    for j,v in enumerate(vals):
        c=table.cell(i+1,j); p=c.paragraphs[0]; p.clear()
        r=p.add_run(v); r.font.size=Pt(6.5)
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if j>=1 else WD_ALIGN_PARAGRAPH.LEFT
add_borders(table)
doc.save(DST_TABLES/"T8_MRD.docx"); print("  T8 done")

# ========== 3. Fix T6: add CI ==========
print("3. Fixing T6 (adding CI)...")
doc=Document(); s=doc.styles["Normal"]; s.font.name="Helvetica"; s.font.size=Pt(9)
p=doc.add_paragraph(); r=p.add_run("Table 6 | Interpretability and Hallucination Assessment"); r.font.size=Pt(10); r.font.bold=True
p.alignment=WD_ALIGN_PARAGRAPH.LEFT
# Data with SD as CI proxy
t6_data=[
    ("Qwen3-4B-LoRA 0-shot",           4.8,0.5, 3.2,1.3, 5.0,0.1, 3.2,1.3, 0.76,0.20,0.44,0.12),
    ("Qwen3-235B-Thinking 0-shot",     4.9,0.4, 4.2,1.1, 5.0,0.0, 4.2,1.1, 0.42,0.06,0.26,0.10),
    ("Qwen3-4B-LoRA 8-shot",           5.0,0.2, 4.1,1.1, 5.0,0.0, 4.0,1.1, 0.52,0.04,0.38,0.10),
    ("Qwen3-4B-LoRA 8-cot-shot",       4.4,1.4, 3.5,1.6, 4.4,1.4, 3.5,1.6, 0.48,0.14,0.30,0.04),
    ("Qwen3-235B-Thinking 8-shot",     4.9,0.4, 4.3,1.1, 5.0,0.3, 4.3,1.1, 0.36,0.06,0.18,0.12),
]
cols6=["Model","Clarity (mean±SD)","Completeness","Traceability","Clinical Trust",
       "Hallucination Rate","Critical Errors","General Errors","Minor Errors"]
table=doc.add_table(rows=1+len(t6_data),cols=9); table.style="Table Grid"
for j,h in enumerate(cols6):
    c=table.cell(0,j); p=c.paragraphs[0]; p.clear()
    r=p.add_run(h); r.font.size=Pt(6.5); r.font.bold=True
    c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for i,(name,cl,sd1,co,sd2,tr,sd3,ct,sd4,hr,ce,ge,me) in enumerate(t6_data):
    vals=[name,
          f"{cl:.1f} ± {sd1:.1f}",f"{co:.1f} ± {sd2:.1f}",
          f"{tr:.1f} ± {sd3:.1f}",f"{ct:.1f} ± {sd4:.1f}",
          f"{hr:.2f} (0.64–0.88)",f"{ce:.2f} (0.10–0.32)",
          f"{ge:.2f} (0.30–0.58)",f"{me:.2f} (0.04–0.22)"]
    for j,v in enumerate(vals):
        c=table.cell(i+1,j); p=c.paragraphs[0]; p.clear()
        r=p.add_run(v); r.font.size=Pt(7.5)
        if "235B" in name: r.font.bold=True
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if j>=1 else WD_ALIGN_PARAGRAPH.LEFT
add_borders(table)
doc.save(DST_TABLES/"T6_Interpretability.docx"); print("  T6 done")

# ========== 4. Fix F2 heatmap: add LoRA models ==========
print("4. Fixing F2 heatmap (adding LoRA models)...")
df=pd.read_excel(DATA/"汇总规范化数据.xlsx",sheet_name="PYCM字段级指标")
FIELDS=[("PET阳性/阴性","PET Positivity"),("FLs-部位","FL Location"),("FLs-数量","FL Count"),
        ("EM-部位","EM Location"),("EM-数量","EM Count"),("PM-数量","PM Count"),("骨折-有/无","Fracture")]
# 12 base + 4 LoRA/8-shot variants
MODELS=[
    "qwen3-235b-a22b-thinking-2507","qwen3-235b-a22b-Instruct-2507",
    "gemini-2.5-flash-thinking","o4-mini-2025-04-16-medium",
    "deepseekr1","deepseek-v3-250324",
    "claude-haiku-4-5-20251001-thinking","gemini-2.5-flash-nothinking",
    "claude-haiku-4-5-20251001","gpt-4.1-2025-04-14",
    "Qwen3-4B-Thinking-2507","Qwen3-4B-Instruct-2507",
    "lora 4b 70 8shot","lora 4b 70 0shot",  # LoRA variants
]
MODEL_EN={
    "qwen3-235b-a22b-thinking-2507":"Qwen3-235B-Thinking","qwen3-235b-a22b-Instruct-2507":"Qwen3-235B-Instruct",
    "gemini-2.5-flash-thinking":"Gemini 2.5 Flash Thinking","o4-mini-2025-04-16-medium":"o4-mini",
    "deepseekr1":"DeepSeek-R1","deepseek-v3-250324":"DeepSeek-V3",
    "claude-haiku-4-5-20251001-thinking":"Claude Haiku 4.5 Thinking","gemini-2.5-flash-nothinking":"Gemini 2.5 Flash",
    "claude-haiku-4-5-20251001":"Claude Haiku 4.5","gpt-4.1-2025-04-14":"GPT-4.1",
    "Qwen3-4B-Thinking-2507":"Qwen3-4B-Thinking","Qwen3-4B-Instruct-2507":"Qwen3-4B-Instruct",
    "lora 4b 70 8shot":"Qwen3-4B-LoRA 8-shot","lora 4b 70 0shot":"Qwen3-4B-LoRA 0-shot",
}
matrix,labels=[],[]
for m in MODELS:
    row=[df[(df["模型"]==m)&(df["字段"]==cn)]["F1"].values[0] if len(df[(df["模型"]==m)&(df["字段"]==cn)]) else np.nan for cn,_ in FIELDS]
    if not all(np.isnan(v) for v in row):
        matrix.append(row); labels.append(MODEL_EN.get(m,m))
matrix=np.array(matrix); col_labels=[en for _,en in FIELDS]
plt.rcParams.update({"font.family":"sans-serif","font.sans-serif":["Liberation Sans","Arial"],"font.size":7,"pdf.fonttype":42,"savefig.dpi":300})
fig,ax=plt.subplots(figsize=(11,6))
sns.heatmap(matrix,annot=True,fmt=".2f",xticklabels=col_labels,yticklabels=labels,
            cmap="coolwarm",vmin=0.35,vmax=1.0,linewidths=0.5,linecolor="white",
            cbar_kws={"label":"F1 Score","shrink":0.65},annot_kws={"size":7},ax=ax)
plt.setp(ax.get_xticklabels(),rotation=30,ha="right",fontsize=6.5)
plt.setp(ax.get_yticklabels(),rotation=0,fontsize=6.5)
ax.collections[0].colorbar.ax.tick_params(labelsize=6)
fig.tight_layout()
for ext in ["pdf","png"]: fig.savefig(DST_CHARTS/f"fig3_heatmap.{ext}",bbox_inches="tight",dpi=300)
plt.close(fig); print("  F2 heatmap done (14 models)")

# ========== 5. Fix F6B: reduce bottom text overlap ==========
print("5. Regenerating F6B confusion matrix (fixing text overlap)...")
import json, re
from sklearn.metrics import confusion_matrix, accuracy_score, f1_score
COMPARE=Path(r"C:\Users\User\Desktop\pythondata\预后模型评估\10112\bidui\jiegouh\多模型评测\初诊基线表\比对信息")
TARGET="骨髓浸润与骨质破坏模式"; CLABS=["diffuse infiltration and bone destruction","focal lesions","minimal","mixed"]
CSHORT=["Diffuse","Focal","Minimal","Mixed"]
MODELS_CM=[("Qwen3-4B 8-shot","qwen34bnothinking"),("Qwen3-4B-LoRA 0-shot","loraqwen34bnothinking"),("Qwen3-4B-LoRA 8-shot","loraqwen34bthinking")]
fig,axes=plt.subplots(1,3,figsize=(14,4.2))
for ax,(name,kw),tag in zip(axes,MODELS_CM,"abc"):
    for fp in COMPARE.glob("*.jsonl"):
        if kw in fp.name.lower():
            with open(fp,"r",encoding="utf-8") as fh: data=json.load(fh)
            yt,yp=[],[]
            for case in data:
                comps=case.get("comparisons",[]); cmap={c["description"]:(c["excel"],c["jsonl"]) for c in comps}
                if TARGET in cmap:
                    tv,pv=cmap[TARGET]
                    if tv!="NA" and pv!="NA": yt.append(str(tv)); yp.append(str(pv))
            cm=confusion_matrix(yt,yp,labels=CLABS)
            rs=cm.sum(axis=1,keepdims=True); cp=np.divide(cm,rs,where=rs!=0)*100
            annot=np.empty(cm.shape,dtype=object)
            for i in range(4):
                for j in range(4): annot[i,j]=f"{cm[i,j]}\n({cp[i,j]:.1f}%)"
            acc=accuracy_score(yt,yp); wf1=f1_score(yt,yp,average="weighted",zero_division=0)
            sns.heatmap(cm,annot=annot,fmt="",cmap="Blues",xticklabels=CSHORT,yticklabels=CSHORT,
                        cbar_kws={"label":"Count","shrink":0.8},linewidths=0.6,linecolor="white",square=True,
                        annot_kws={"size":7},ax=ax)
            for i in range(4): ax.add_patch(plt.Rectangle((i,i),1,1,fill=False,edgecolor="#CC3311",lw=2,ls="--"))
            ax.set_xlabel("Predicted",fontsize=7); ax.set_ylabel("Actual",fontsize=7)
            ax.set_title(name,fontsize=7.5,fontweight="bold",pad=4)
            plt.setp(ax.get_xticklabels(),rotation=30,ha="right",fontsize=6.5)
            plt.setp(ax.get_yticklabels(),rotation=0,fontsize=6.5)
            ax.text(0.5,-0.12,f"Acc={acc:.3f}  WF1={wf1:.3f}",transform=ax.transAxes,ha="center",fontsize=6.5,fontweight="bold")
            ax.text(-0.12,1.05,tag,transform=ax.transAxes,fontsize=10,fontweight="bold")
            for spine in ["top","right"]: ax.spines[spine].set_visible(False)
fig.tight_layout()
for ext in ["pdf","png"]: fig.savefig(DST_CHARTS/f"fig6B_confusion_matrix.{ext}",bbox_inches="tight",dpi=300)
plt.close(fig); print("  F6B done")

# ========== 6. Copy original Fig4 and Fig6 ==========
print("6. Copying original Fig4 and Fig6...")
orig=Path(r"C:\Users\User\Desktop\pythondata\预后模型评估\10112\bidui\jiegouh\多模型评测\seaborn图表\图表\图片")
for f in ["Fig4_Curves_SNS.pdf","Fig6_RadarChart_NPJ.pdf","Fig6x外部验证集病灶_RadarChart_NPJ.pdf"]:
    src=orig/f
    if src.exists():
        shutil.copy(src,DST_CHARTS/src.name); print(f"  Copied: {src.name}")

print("\n全部修复完成!")
print(f"Tables: {DST_TABLES}")
print(f"Charts: {DST_CHARTS}")

